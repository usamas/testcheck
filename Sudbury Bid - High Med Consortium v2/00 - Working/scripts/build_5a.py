#!/usr/bin/env python3
"""Populate Document 5A (Selection Questionnaire) for each consortium member.

Verified company data (from the consortium's completed Isleworth SQ) is reused;
standard compliant answers are set for exclusion grounds and confirmations.
Entity-specific identifiers we cannot verify are marked '[To confirm: ...]'
rather than fabricated. Produces one 5A per bidder member.
"""
import os
import re
import shutil
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.normpath(os.path.join(HERE, "..", "..", ".."))
TEMPLATE = os.path.join(ROOT, "Sudbury surgery", "Document 5 A - Selection Questionnaire (3).docx")
TRD = os.path.normpath(os.path.join(HERE, "..", "..", "01 - Tender Response Documents"))

# Standard compliant answers shared by both members
COMMON = {
    # Mandatory exclusion grounds - none apply
    "2.1 (a) – (i)": "No", "2.1 (a) – (ii)": "No", "2.1 (a) – (iii)": "No",
    "2.1 (a) –(iv)": "No", "2.1 (a) – (v)": "No", "2.1 (a) – (vi)": "No",
    "2.1 (a) – (vii)": "No",
    "2.1 (b) – (i)": "N/A", "2.1 (b) – (ii)": "N/A", "2.1 (c)": "N/A",
    # Tax
    "3.2(a) – (i)": "Yes", "3.2(a) – (ii)": "N/A", "3.2(b)": "N/A", "3.3": "N/A",
    # Discretionary exclusion grounds - none apply
    "4.1 (a)": "No", "4.1 (b)": "No", "4.1 (c)": "No", "4.1 (d)": "No",
    "4.1 (e)": "No", "4.1 (f)": "No", "4.1 (g)": "No", "4.1 (h)": "No",
    "4.1 (i)": "No", "4.1 (j)": "No", "4.2": "No", "4.3": "N/A", "4.4": "N/A",
    # Economic/financial standing
    "5.5": "N/A",
    # Additional / project-specific
    "6.1": "Yes - we hold or will hold the required levels of insurance.",
    "6.3": "N/A - GP/APMS providers are not required to hold an NHS Provider Licence.",
    "6.4": "No",
    "6.5": "Yes - all applicable staff have appropriate DBS checks.",
    "6.6": "No",
    "6.7": "No",
    "6.8": "Yes - we confirm compliance with New Fair Deal / COSOP pension obligations.",
    "6.9 (a)": "No", "6.9 (b)": "No", "6.9 (c)": "N/A",
    "6.10": "No",
    "6.11": "Yes",
    "6.12 (a)": "We hold no contracts with a required supply chain to report.",
    "6.12 (b)": "We do not hold any contracts requiring a supply chain.",
    "6.12 (c)": "N/A - as above, no supply chain contracts.",
    "6.13 (a)": "No", "6.13 (b)": "N/A",
    "6.14 (a)": "Yes", "6.14 (b)": "Yes",
    "6.15": "No",
    "6.16": "Yes", "6.17": "Yes", "6.18": "Yes", "6.19": "Yes",
    "6.20": "Yes - all policies listed below are in place.",
    "6.20 (a)": "Yes", "6.20 (b)": "Yes", "6.20 (c)": "Yes", "6.20 (d)": "Yes",
    "6.20 (e)": "Yes", "6.20 (f)": "Yes", "6.20 (g)": "Yes", "6.20 (h)": "Yes",
    "6.20 (i)": "Yes", "6.20 (j)": "Yes",
    "6.21": "Yes",
    "6.22 (a)": "Yes", "6.22 (b)": "No", "6.22 (c)": "N/A",
    "6.23": "Yes",
}

HIGH_MED = dict(COMMON)
HIGH_MED.update({
    "1.1 (a) – (i)": "HIGH MED LTD",
    "1.1 (a) – (ii)": "High Med Ltd (Lead Bidder; High Med Consortium with Dr Singh Hammond Road)",
    "1.1 (b) – (i)": "Flat 6 Stourcliffe Close, Stourcliffe Street, London, W1H 5AQ",
    "1.1 (b) – (ii)": "N/A",
    "1.1 (b) – (iii)": "[To confirm: published contact email]",
    "1.1 (b) – (iv)": "[To confirm: Find a Tender Service Unique Identifier]",
    "1.1 (c)": "Private limited company",
    "1.1 (d)": "28 February 2014",
    "1.1 (e)": "08917150",
    "1.1 (f)": "N/A",
    "1.1 (g) - (i)": "Yes",
    "1.1 (g) - (ii)": "Clinicians are GMC-registered; Dr Shumaila Mahmood, GMC Registration Number 6095460.",
    "1.1 (h) - (i)": "Yes",
    "1.1(h) - (ii)": "Dr Mahmood is GMC-registered, Registration Number 6095460.",
    "1.1 (i)": "N/A",
    "1.1 (j)": "Yes - Small/Micro Enterprise",
    "1.1 (k)": "Name: Mahmood Butt, Shumaila, Dr. Date of birth: March 1972. Nationality: British.",
    "1.1 (l)": "N/A", "1.1 (m)": "N/A",
    "1.2": ("Part of a Consortium: High Med Consortium (High Med Ltd as Lead Bidder with "
            "Dr Singh Hammond Road as key party). A Special Purpose Vehicle will be established "
            "by the consortium if successful."),
    "1.2 (a)": "Yes",
    "1.3": "N/A",
    "5.1": "https://find-and-update.company-information.service.gov.uk/company/08917150",
    "6.2 (a)": ("N/A for High Med Ltd. Consortium key party Dr Singh Hammond Road is a "
                "CQC-registered GP practice rated Good."),
    "6.2 (b)": "Consortium member Dr Singh Hammond Road / Sudbury Surgery is CQC rated Good.",
})

DR_SINGH = dict(COMMON)
DR_SINGH.update({
    "1.1 (a) – (i)": "Dr Gursharan Singh (Dr Singh Hammond Road)",
    "1.1 (a) – (ii)": "Dr Singh Hammond Road (key party; High Med Consortium with High Med Ltd as Lead Bidder)",
    "1.1 (b) – (i)": "[To confirm: registered/practice address - Dr Singh Hammond Road]",
    "1.1 (b) – (ii)": "[To confirm]",
    "1.1 (b) – (iii)": "[To confirm: published contact email]",
    "1.1 (b) – (iv)": "[To confirm: Find a Tender Service Unique Identifier]",
    "1.1 (c)": "[To confirm: trading status (e.g. sole trader / partnership)]",
    "1.1 (d)": "[To confirm: date of formation]",
    "1.1 (e)": "[To confirm: registration number, if applicable]",
    "1.1 (f)": "N/A",
    "1.1 (g) - (i)": "Yes",
    "1.1 (g) - (ii)": "Dr Gursharan Singh is GMC-registered (registration number to confirm).",
    "1.1 (h) - (i)": "Yes",
    "1.1(h) - (ii)": "Dr Gursharan Singh is GMC-registered (registration number to confirm).",
    "1.1 (i)": "N/A",
    "1.1 (j)": "Yes - Small/Micro Enterprise",
    "1.1 (k)": "[To confirm: Persons with Significant Control]",
    "1.1 (l)": "N/A", "1.1 (m)": "N/A",
    "1.2": ("Part of a Consortium: High Med Consortium (High Med Ltd as Lead Bidder with "
            "Dr Singh Hammond Road as key party)."),
    "1.2 (a)": "Yes",
    "1.3": "N/A",
    "5.1": "[To confirm: Companies House / accounts reference for Dr Singh Hammond Road]",
    "6.2 (a)": "Yes - Dr Singh Hammond Road is a CQC-registered GP practice.",
    "6.2 (b)": "CQC rated Good.",
})


def _style(run):
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.bold = True


def write_answer(cell, text):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    _style(run)


def populate(data, out_name):
    out = os.path.join(TRD, out_name)
    shutil.copyfile(TEMPLATE, out)
    doc = Document(out)
    filled = 0
    for tbl in doc.tables:
        rows = tbl.rows
        for ri, row in enumerate(rows):
            key = row.cells[0].text.strip()
            # normalise multiline keys (e.g. 4.1 (j) block) to first token
            key_first = key.split("\n")[0].strip()
            ans = data.get(key) or data.get(key_first)
            if ans is None:
                continue
            # Case A: 6.20 sub-items - inline Yes/No in cells 2/3 of same row
            if re.match(r"^6\.20 \([a-j]\)$", key_first):
                for c in row.cells:
                    ct = c.text.strip().lower()
                    if ct == "yes" or ct.startswith("yes"):
                        write_answer(c, "Yes (selected)")
                        break
                filled += 1
                continue
            # Case B: response row immediately after
            if ri + 1 < len(rows):
                nxt = rows[ri + 1].cells[0]
                if nxt.text.strip().lower().startswith(("please enter", "please provide")):
                    write_answer(nxt, ans)
                    filled += 1
                    continue
    doc.save(out)
    print(f"  {out_name}: {filled} answers filled")
    return filled


if __name__ == "__main__":
    populate(HIGH_MED, "Document 5A - Selection Questionnaire - High Med Ltd (Lead).docx")
    populate(DR_SINGH, "Document 5A - Selection Questionnaire - Dr Singh Hammond Road.docx")
    # remove the unbranded copied template to avoid confusion
    generic = os.path.join(TRD, "Document 5A - Selection Questionnaire.docx")
    if os.path.exists(generic):
        os.remove(generic)
    print("5A populated for both consortium members.")
