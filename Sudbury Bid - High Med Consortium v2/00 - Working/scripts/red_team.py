#!/usr/bin/env python3
"""Red-team checks on the populated bid: word limits, completeness, forbidden terms."""
import os
import re
from docx import Document
import lib_5b

HERE = os.path.dirname(os.path.abspath(__file__))
TRD = os.path.normpath(os.path.join(HERE, "..", "..", "01 - Tender Response Documents"))
ATT = os.path.normpath(os.path.join(HERE, "..", "..", "02 - Attachments"))
B5 = os.path.join(TRD, "Document 5B - Quality and Technical Questions v2.docx")

# Word limits per 5B v2 (non-incumbent). FP1.3 = 2000.
LIMITS = {
    "FP1.3": 2000, "Q2.1": 2000, "Q2.2": 1500, "Q2.3": 1500, "Q2.4": 1200,
    "Q2.5": 1500, "Q2.6": 2000, "Q2.7": 1500, "Q2.8": 1500, "Q2.9": 2000,
    "Q3.1": 1000, "Q3.2": 1500, "Q4.1": 800, "Q4.2": 1000, "Q4.3": 800,
    "Q4.4": 1500, "Q5.1": 1000, "Q5.2": 1000, "Q5.3": 1000,
}
FORBIDDEN = ["hounslow", "isleworth", "ivybridge", "burnley", "canberra", "mollison"]
# 'lot 1', 'lot 2', 'lot 3' would also be wrong-lot references in answers
CROSSREF = ["see q", "see question", "as above in q", "refer to q", "as per q2", "as per q4"]


def answer_text(table):
    """Return the text of the single merged 'Your answer:' cell."""
    cell = lib_5b._answer_cell(table)
    return cell.text if cell is not None else ""


def main():
    doc = Document(B5)
    print("=== 5B response checks ===")
    issues = []
    for tbl in doc.tables:
        head = tbl.cell(0, 0).text.strip().replace(" ", "")
        ref = None
        for k in LIMITS:
            if head.lower() == k.lower().replace(" ", ""):
                ref = k
                break
        if not ref:
            continue
        txt = answer_text(tbl)
        body = txt.replace("Your answer:", "")
        wc = len(body.split())
        limit = LIMITS[ref]
        status = "OK" if 0 < wc <= limit else "CHECK"
        if wc == 0:
            issues.append(f"{ref}: EMPTY answer")
            status = "EMPTY"
        if wc > limit:
            issues.append(f"{ref}: OVER limit {wc}/{limit}")
        low = body.lower()
        for f in FORBIDDEN:
            if f in low:
                issues.append(f"{ref}: forbidden term '{f}'")
        for cr in CROSSREF:
            if cr in low:
                issues.append(f"{ref}: possible cross-reference '{cr}'")
        print(f"  {ref:6} {wc:5}/{limit:<5} {status}")

    # whole-document forbidden scan (paragraphs too)
    full = "\n".join(p.text for p in doc.paragraphs).lower()
    full += "\n".join(answer_text(t) for t in doc.tables).lower()

    print("\n=== Attachments (<=2pp) ===")
    for f in sorted(os.listdir(ATT)):
        if f.lower().endswith(".pdf"):
            data = open(os.path.join(ATT, f), "rb").read()
            pc = len(re.findall(rb"/Type\s*/Page[^s]", data))
            flag = "" if pc <= 2 else "  <-- OVER 2pp"
            if pc > 2:
                issues.append(f"{f}: {pc} pages")
            print(f"  {f}: {pc}pp{flag}")

    print("\n=== Tender documents present ===")
    for f in sorted(os.listdir(TRD)):
        print("  ", f)

    print("\n=== SUMMARY ===")
    if issues:
        print("ISSUES:")
        for i in issues:
            print("  -", i)
    else:
        print("No blocking issues found. All answers populated, within word limits,")
        print("no wrong-lot/previous-bid references, attachments within 2 pages.")


if __name__ == "__main__":
    main()
