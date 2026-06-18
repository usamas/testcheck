"""Shared helpers for populating Document 5B (Lot 4 Sudbury Surgery v2).

Answer content is expressed as a list of "blocks":
  ("h", "Heading")            -> bold heading paragraph
  ("p", "paragraph text")     -> normal paragraph
  ("b", ["item", ...])        -> bullet list
  ("t", [[h1,h2], [a,b], ...]) -> table, first row treated as bold header

The populate routine locates each question table by its reference label
(e.g. "Q 2.1" / "R 2.1"), writes the rendered blocks into the "Your answer:"
cell, fills the "Number of words used:" cell, and marks attachment rows "Yes".
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ANSWER_FONT = "Arial"
ANSWER_SIZE = Pt(10)


def _count_words(blocks):
    n = 0
    for kind, payload in blocks:
        if kind in ("h", "p"):
            n += len(payload.split())
        elif kind == "b":
            for item in payload:
                n += len(item.split())
        elif kind == "t":
            for row in payload:
                for cell in row:
                    n += len(str(cell).split())
    return n


def _style_run(run, bold=False):
    run.font.name = ANSWER_FONT
    run.font.size = ANSWER_SIZE
    run.bold = bold


def _add_paragraph(cell, text, bold=False, bullet=False):
    p = cell.add_paragraph()
    if bullet:
        p.paragraph_format.left_indent = Pt(12)
        run = p.add_run("\u2022  " + text)
    else:
        run = p.add_run(text)
    _style_run(run, bold=bold)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_table(cell, rows):
    cols = max(len(r) for r in rows)
    tbl = cell.add_table(rows=len(rows), cols=cols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            val = str(row[ci]) if ci < len(row) else ""
            tc = tbl.cell(ri, ci)
            tc.text = ""
            p = tc.paragraphs[0]
            run = p.add_run(val)
            _style_run(run, bold=(ri == 0))
    return tbl


def render_blocks(cell, blocks):
    """Append rendered blocks to a table cell (after its existing label)."""
    for kind, payload in blocks:
        if kind == "h":
            _add_paragraph(cell, payload, bold=True)
        elif kind == "p":
            _add_paragraph(cell, payload)
        elif kind == "b":
            for item in payload:
                _add_paragraph(cell, item, bullet=True)
        elif kind == "t":
            _add_table(cell, payload)


def _find_label_cell(table, label_regex):
    rx = re.compile(label_regex)
    for ri, row in enumerate(table.rows):
        for ci, c in enumerate(row.cells):
            if rx.match(c.text.strip()):
                return ri, ci
    return None, None


def _answer_cell(table):
    """Return the cell that contains the 'Your answer:' label."""
    for row in table.rows:
        for c in row.cells:
            if c.text.strip().lower().startswith("your answer"):
                return c
    return None


def _set_word_count(table, count):
    for ri, row in enumerate(table.rows):
        for ci, c in enumerate(row.cells):
            if c.text.strip().lower().startswith("number of words used"):
                # the adjacent cell to the right holds the figure
                cells = row.cells
                target = None
                for cc in cells:
                    if not cc.text.strip():
                        target = cc
                        break
                if target is not None:
                    target.text = ""
                    run = target.paragraphs[0].add_run(str(count))
                    _style_run(run)
                return True
    return False


def _mark_attachments_yes(table, attach_labels):
    """For rows whose first cell matches an attachment label, set Yes."""
    for row in table.rows:
        first = row.cells[0].text.strip()
        for lbl in attach_labels:
            if first.lower().startswith(lbl.lower()):
                for c in row.cells:
                    if "yes" in c.text.strip().lower() and "no" in c.text.strip().lower():
                        c.text = ""
                        run = c.paragraphs[0].add_run("Yes")
                        _style_run(run, bold=True)
                        break


def find_question_table(doc, ref):
    """Find the table for a question by its full reference label.

    ref examples: 'FP1.3', 'Q2.1', 'Q4.4'. Matches a table whose first cell
    normalises (lowercase, no spaces) exactly to the ref.
    """
    want = ref.strip().lower().replace(" ", "")
    for tbl in doc.tables:
        head = tbl.cell(0, 0).text.strip().lower().replace(" ", "")
        if head == want:
            return tbl
    return None


def mark_yes_no_after(doc, ref, choice="Yes"):
    """Find the question table for ref, then mark the following Yes/No table."""
    tables = doc.tables
    qtbl = find_question_table(doc, ref)
    if qtbl is None:
        raise RuntimeError(f"Question table not found for ref {ref!r} (yes/no)")
    qidx = None
    for i, t in enumerate(tables):
        if t._tbl is qtbl._tbl:
            qidx = i
            break
    if qidx is None:
        raise RuntimeError(f"Could not index question table for ref {ref!r}")
    for t in tables[qidx + 1:]:
        cells0 = [c.text.strip().lower() for c in t.rows[0].cells]
        if "yes" in cells0 and "no" in cells0:
            for c in t.rows[0].cells:
                if c.text.strip().lower() == choice.lower():
                    c.text = ""
                    run = c.paragraphs[0].add_run(f"\u2612 {choice}")
                    _style_run(run, bold=True)
                    return True
    raise RuntimeError(f"No Yes/No table found after ref {ref!r}")


def populate_answer(doc, ref, blocks, attachments=None, verbose=True):
    tbl = find_question_table(doc, ref)
    if tbl is None:
        raise RuntimeError(f"Question table not found for ref {ref!r}")
    cell = _answer_cell(tbl)
    if cell is None:
        raise RuntimeError(f"'Your answer:' cell not found for ref {ref!r}")
    render_blocks(cell, blocks)
    wc = _count_words(blocks)
    _set_word_count(tbl, wc)
    if attachments:
        _mark_attachments_yes(tbl, attachments)
    if verbose:
        print(f"  {ref}: {wc} words"
              + (f", attachments: {', '.join(attachments)}" if attachments else ""))
    return wc
