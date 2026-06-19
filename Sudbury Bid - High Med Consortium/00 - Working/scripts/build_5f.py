#!/usr/bin/env python3
"""Populate Document 5F (Form of Offer and Declarations) for the consortium."""
import os
import shutil
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.normpath(os.path.join(HERE, "..", "..", ".."))
TEMPLATE = os.path.join(ROOT, "Sudbury surgery", "Document 5 F - Form of Offer and Declarations.docx")
TRD = os.path.normpath(os.path.join(HERE, "..", "..", "01 - Tender Response Documents"))
OUT = os.path.join(TRD, "Document 5F - Form of Offer and Declarations.docx")


def setcell(cell, text):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)


def main():
    shutil.copyfile(TEMPLATE, OUT)
    doc = Document(OUT)

    # T0 - signature block (authorised representative of the lead bidder)
    sig = {
        "Name (print):": "Dr Shumaila Mahmood",
        "Signature:": "[To be signed on submission]",
        "Position / Job Title:": "Director / Clinical Director",
        "For and on behalf of:": "High Med Ltd (Lead Bidder, High Med Consortium) - Lot 4 Sudbury Surgery",
        "Date:": "[Date of submission]",
    }
    t0 = doc.tables[0]
    for row in t0.rows:
        label = row.cells[0].text.strip()
        if label in sig:
            setcell(row.cells[1], sig[label])

    # T1 - organisation declaration of interest
    t1 = doc.tables[1]
    org_decl = {
        "Name of Organisation:": "High Med Consortium (High Med Ltd, Lead Bidder + Dr Singh Hammond Road, key party)",
        "Details of interests held:": ("Dr Gursharan Singh (consortium key party) is an existing "
                                       "GP practitioner in the Brent area. Declared for transparency; "
                                       "no conflict of interest is considered to arise in relation to "
                                       "this procurement."),
    }
    for row in t1.rows:
        label = row.cells[0].text.strip()
        if label in org_decl:
            setcell(row.cells[1], org_decl[label])
        elif label.startswith("Provision of Goods") or label.startswith("Any other connection"):
            setcell(row.cells[1], "None")

    # T2 - relevant person declaration
    t2 = doc.tables[2]
    for row in t2.rows:
        label = row.cells[0].text.strip()
        if label.startswith("Name and Role Title"):
            setcell(row.cells[1], "Dr Gursharan Singh - consortium key party / clinical lead")
        elif label.startswith("Details of interests held"):
            setcell(row.cells[1], "Existing local GP practitioner; declared for transparency.")
        elif label.startswith("Provision of Goods") or label.startswith("Any other connection"):
            setcell(row.cells[1], "None")

    doc.save(OUT)
    print("5F populated: signature block + declarations of interest.")


if __name__ == "__main__":
    main()
